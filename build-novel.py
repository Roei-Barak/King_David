#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""בונה רומן פרוזה אחד (Word) מפרקי הפרוזה שב-prose-chapters/."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/home/user/King_David")
CH = ROOT / "prose-chapters"
FONT = "David"          # גופן עברי קלאסי; אם אין — Word ימיר ל-Arial
FONT_FALLBACK = "Arial"

# חלוקת מערכות לפי מספר פרק
PART_BREAKS = {
    1:  ("פֵּתַח", None),
    2:  ("חֵלֶק רִאשׁוֹן", "מִן הָרוֹעֶה אֶל הַמֶּלֶךְ"),
    14: ("חֵלֶק שֵׁנִי", "הַמֶּלֶךְ"),
}

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1.1); s.bottom_margin = Inches(1.1)
    s.left_margin = Inches(1.3); s.right_margin = Inches(1.3)

style = doc.styles["Normal"]
style.font.name = FONT_FALLBACK
style.font.size = Pt(12.5)
style._element.rPr.rFonts.set(qn("w:cs"), FONT_FALLBACK)
style.paragraph_format.space_after = Pt(0)

def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def par(text, size=12.5, bold=False, italic=False, color=None, center=False,
        sb=0, sa=6, ls=1.5, justify=True, first_indent=0.0):
    if text is None: return
    p = doc.add_paragraph(); set_rtl(p)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = Pt(size * ls)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    run = p.add_run(text)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.name = FONT_FALLBACK
    run.font._element.rPr.rFonts.set(qn("w:cs"), FONT_FALLBACK)
    if color: run.font.color.rgb = color
    return p

def poem(text, size=12):
    p = doc.add_paragraph(); set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(size*1.6)
    run = p.add_run(text)
    run.font.size = Pt(size); run.italic = True
    run.font.name = FONT_FALLBACK
    run.font._element.rPr.rFonts.set(qn("w:cs"), FONT_FALLBACK)
    run.font.color.rgb = RGBColor(0x2a,0x2a,0x55)

def clean(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    return s.strip()

# ---- שער ----
for _ in range(5): doc.add_paragraph()
par("מִזְמוֹר לְדָוִד", size=40, bold=True, center=True, sa=12, justify=False)
par("רוֹמָן", size=20, italic=True, center=True, sa=8, justify=False)
par("סִפּוּר חַיָּיו שֶׁל דָּוִד הַמֶּלֶךְ — אֱמוּנָה דֶּרֶךְ יְסוּרִים",
    size=13, center=True, justify=False, color=RGBColor(0x55,0x55,0x55))
doc.add_page_break()

def render_chapter(md):
    lines = md.split("\n")
    in_poem = False
    poem_buf = []
    def flush_poem():
        nonlocal poem_buf, in_poem
        if poem_buf:
            doc.add_paragraph()
            for l in poem_buf: poem(l)
            doc.add_paragraph()
        poem_buf = []; in_poem = False
    for raw in lines:
        line = raw.rstrip()
        # כותרת פרק
        m = re.match(r"^#\s+(.*)$", line)
        if m:
            flush_poem()
            doc.add_page_break()
            for _ in range(2): doc.add_paragraph()
            par(clean(m.group(1)), size=22, bold=True, center=True, sa=18, justify=False)
            continue
        m2 = re.match(r"^#{2,6}\s+(.*)$", line)
        if m2:
            flush_poem()
            par(clean(m2.group(1)), size=14, bold=True, sb=10, sa=4, justify=False,
                color=RGBColor(0x33,0x33,0x55))
            continue
        # מפריד קטעים
        if re.match(r"^\s*(---+|\*\s*\*\s*\*|⸻)\s*$", line):
            flush_poem()
            par("✦", center=True, size=12, color=RGBColor(0x99,0x99,0x99), sb=4, sa=4, justify=False)
            continue
        # ציטוט / שיר
        if line.strip().startswith(">"):
            t = clean(line.strip()[1:].strip())
            if t:
                poem_buf.append(t); in_poem = True
            continue
        # שורה ריקה
        if not line.strip():
            if in_poem: continue   # שמור רצף שיר
            continue
        # שורת פרוזה רגילה
        flush_poem()
        par(clean(line), first_indent=0.3)
    flush_poem()

files = sorted(CH.glob("[0-9][0-9]-*.md"))
print(f"פרקים שנמצאו: {len(files)}")
for f in files:
    num = int(f.name[:2])
    if num in PART_BREAKS:
        title, sub = PART_BREAKS[num]
        doc.add_page_break()
        for _ in range(6): doc.add_paragraph()
        par(title, size=26, bold=True, center=True, sa=8, justify=False)
        if sub: par(sub, size=16, italic=True, center=True, justify=False,
                    color=RGBColor(0x55,0x55,0x55))
    print(f"  + {f.name}")
    render_chapter(f.read_text(encoding="utf-8"))

out = ROOT / "מזמור-לדוד-רומן.docx"
doc.save(out)
print(f"\nנשמר: {out}")
print(f"פסקאות: {len(doc.paragraphs)}")
