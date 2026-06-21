#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""בונה מסמך Word אחד מכל הסצנות, מסודר כספר לקריאה והערות."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/home/user/King_David")
SCENES = ROOT / "scenes"

BOOK = [
    ("חלק א'", "מַעֲרָכָה א — מִן הָרוֹעֶה לַמֶּלֶךְ", None),
    (None, None, "prologue.md"),
    (None, None, "scene-01-05-full.md"),
    ("הרחבות", "הרחבות המדבר (תמונה 4) — לפי סדר מקראי", None),
    (None, None, "scene-jonathan-farewell.md"),
    (None, None, "scene-gath-madness.md"),
    (None, None, "scene-nob-massacre.md"),
    (None, None, "scene-04-the-cave.md"),
    (None, None, "scene-04b-wilderness-additions.md"),
    (None, None, "scene-nabal-abigail.md"),
    (None, None, "scene-cave-hill.md"),
    (None, None, "scene-05b-witch-endor.md"),
    (None, None, "scene-ziklag.md"),
    (None, None, "scene-philistines-reject.md"),
    (None, None, "scene-gilboa.md"),
    ("חלק ב'", "מַעֲרָכָה ב — הַמֶּלֶךְ", None),
    (None, None, "scene-06-lament.md"),
    (None, None, "scene-07-brothers-war.md"),
    (None, None, "scene-ishbosheth.md"),
    (None, None, "scene-08-dance.md"),
    (None, None, "scene-09-not-you.md"),
    (None, None, "scene-chr25-singers.md"),
    (None, None, "scene-10-mephibosheth.md"),
    (None, None, "scene-10b-ammon-war.md"),
    (None, None, "scene-bathsheba-uriah.md"),
    (None, None, "bathhouse.md"),
    (None, None, "scene-13b-amnon-tamar.md"),
    (None, None, "scene-13e-absalom-return.md"),
    (None, None, "scene-13c-revolt.md"),
    (None, None, "scene-13d-absalom.md"),
    (None, None, "scene-13f-david-return.md"),
    (None, None, "scene-13g-rizpah.md"),
    (None, None, "scene-water-libation.md"),
    (None, None, "scene-14-hallelujah.md"),
    (None, None, "scene-22-23-david-song.md"),
    (None, None, "scene-15-testament.md"),
    ("נספח", "נספח — תת-עלילות", None),
    (None, None, "scene-tzruya-david.md"),
]

FONT = "Arial"

doc = Document()

# שוליים צרים יותר
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# סגנון בסיס — Arial, RTL
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:cs"), FONT)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(3)

def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    # bidi — כיוון RTL לפסקה
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_par(text, size=11, bold=False, italic=False, color=None,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=3, line_spacing=1.15):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * line_spacing * 1.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run.font._element.rPr.rFonts.set(qn("w:cs"), FONT)
    if color:
        run.font.color.rgb = color
    return p

# --- עמוד שער ---
for _ in range(3):
    doc.add_paragraph()
add_par("מִזְמוֹר לְדָוִד", size=36, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_par("הַצַּדִּיק שֶׁכּוֹתֵב שִׁירִים", size=18, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
add_par("מחזמר במה על חיי דוד המלך", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_par("טיוטה לקריאה והערות", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80,0x80,0x80))
doc.add_page_break()

# --- הערת קריאה ---
add_par("לקורא", size=16, bold=True, space_after=6)
add_par("מסמך זה מאחד את כל הסצנות של המחזמר בסדר קריאה רציף, כספר.", space_after=3)
add_par("שלוש השפות: דיאלוג = פרוזה מליצית מקראית · מונולוג פנימי = שיר חופשי על דרך תהילים · שיר מוזיקלי = חרוז ומטר קבועים.", space_after=3)
add_par("סימון [להכריע: רסיטטיב/פזמון] — מקום שטרם הוכרע אם להלחינו כרסיטטיב או לעבדו לפזמון.", space_after=3)
add_par("הערות *(כך)* הן הוראות במאי. ציטוטים מקראיים מסומנים *(ספר פרק, פסוק)*.", space_after=3)
doc.add_page_break()

# --- עיבוד markdown לפסקאות Word ---
def render_md(text):
    lines = text.split("\n")
    in_code = False
    prev_empty = False
    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        # שורה ריקה — מדלגים על רצפים של שורות ריקות
        if not line.strip():
            if not prev_empty:
                prev_empty = True
            continue
        prev_empty = False
        if re.match(r"^\s*\|[-\s:|]+\|\s*$", line):
            continue
        # כותרות
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            txt = clean_inline(m.group(2))
            size = {1:18,2:15,3:13,4:12,5:11,6:11}.get(level,11)
            sa = 5 if level <= 2 else 3
            add_par(txt, size=size, bold=True, space_after=sa)
            continue
        # ציטוט/בלוק
        if line.strip().startswith(">"):
            txt = clean_inline(line.strip()[1:].strip())
            add_par(txt, italic=True, color=RGBColor(0x55,0x55,0x55), space_after=2)
            continue
        # שורת טבלה
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            txt = "  —  ".join(c for c in cells if c)
            add_par(clean_inline(txt), size=10, space_after=2)
            continue
        # רשימה
        if re.match(r"^\s*[-*]\s+", line):
            txt = clean_inline(re.sub(r"^\s*[-*]\s+","",line))
            add_par("•  " + txt, space_after=2)
            continue
        # מפריד
        if re.match(r"^\s*---+\s*$", line):
            add_par("⸻", align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xAA,0xAA,0xAA), space_after=3)
            continue
        # רגיל
        add_par(clean_inline(line), space_after=3)

def clean_inline(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    return s

# --- בניית הספר ---
for part, title, fname in BOOK:
    if fname is None:
        doc.add_page_break()
        add_par(part, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                color=RGBColor(0x55,0x55,0x55), space_after=4)
        add_par(title, size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        continue
    fpath = SCENES / fname
    if not fpath.exists():
        continue
    doc.add_page_break()
    render_md(fpath.read_text(encoding="utf-8"))

COVERAGE = [
    ("פרולוג (ישעיה טז, מזמור עח)", "פרולוג", "חלק א'"),
    ("שמ\"א א–ה (שמואל, גוליית, חצר שאול)", "נעורים ועלייה", "חלק א'"),
    ("שמ\"א כג — פרידת יונתן", "פרידת יונתן", "חלק א'"),
    ("שמ\"א כא — דוד בגת", "שיגעון בגת", "חלק א'"),
    ("שמ\"א כב — טבח כהני נוב", "טבח נוב", "חלק א'"),
    ("שמ\"א כד — מערת עין גדי", "המערה", "חלק א'"),
    ("שמ\"א כג–כו — נדודי המדבר", "הוספות המדבר", "חלק א'"),
    ("שמ\"א כה — נבל ואביגיל", "נבל ואביגיל", "חלק א'"),
    ("שמ\"א כד/כו — מערה וגבעה", "המערה והגבעה", "חלק א'"),
    ("שמ\"א כח — אוב עין דור", "אשת עין דור", "חלק א'"),
    ("שמ\"א כז, ל — צקלג", "צקלג", "חלק א'"),
    ("שמ\"א כט — פלשתים דוחים", "הפלשתים דוחים", "חלק א'"),
    ("שמ\"א לא — גלבוע", "גלבוע", "חלק א'"),
    ("שמ\"ב א — קינת דוד", "קינה", "חלק ב'"),
    ("שמ\"ב ב–ג — מלחמת האחים", "מלחמת האחים", "חלק ב'"),
    ("שמ\"ב ד — מות איש-בשת", "מות איש-בשת", "חלק ב'"),
    ("שמ\"ב ו — ריקוד לפני הארון", "הריקוד ומיכל", "חלק ב'"),
    ("שמ\"ב ז — לא אתה תבנה", "לא אתה תבנה", "חלק ב'"),
    ("דה\"א כה — 288 המשוררים", "המשוררים", "חלק ב'"),
    ("שמ\"ב ט — מפיבושת", "מפיבושת", "חלק ב'"),
    ("שמ\"ב י — מלחמת עמון", "מלחמת עמון", "חלק ב'"),
    ("שמ\"ב יא–יב — בת-שבע ואוריה", "בת-שבע ואוריה", "חלק ב'"),
    ("שמ\"ב יא — בית המרחץ (מדרש)", "בית המרחץ", "חלק ב'"),
    ("שמ\"ב יג — אמנון ותמר", "אמנון ותמר", "חלק ב'"),
    ("שמ\"ב יד — שיבת אבשלום", "שיבת אבשלום", "חלק ב'"),
    ("שמ\"ב טו–יז — מרד אבשלום", "מרד אבשלום", "חלק ב'"),
    ("שמ\"ב יח — מות אבשלום", "מות אבשלום", "חלק ב'"),
    ("שמ\"ב יט — שיבת דוד", "שיבת דוד", "חלק ב'"),
    ("שמ\"ב כא — רצפה בת איה", "רצפה", "חלק ב'"),
    ("שמ\"ב כג — שפיכת המים", "שפיכת המים", "חלק ב'"),
    ("שמ\"ב כב — שיר דוד", "הללויה", "חלק ב'"),
    ("שמ\"ב כב–כג — שיר ומגיבורים", "שיר דוד והגיבורים", "חלק ב'"),
    ("מל\"א א–ב — הצוואה והמלכת שלמה", "הצוואה", "חלק ב'"),
    ("תת-עלילה — בני צרויה", "בני צרויה", "נספח"),
]

MODERN_REFS = [
    ("שורש המלוכה / פרולוג", "המילטון: \"Alexander Hamilton\" · ליאון מלך: \"Circle of Life\" · לה מיזרבל: פרולוג"),
    ("נעורים ברועה / עלייה", "המילטון: \"My Shot\" / \"Young, Scrappy and Hungry\" · טרובדור: עלייה מהאפר"),
    ("דוד ושאול — קנאה ורדיפה", "המלט: קלאודיוס ורוח אביו · ריצ'רד ג' · ליאון מלך: סקר · מקבת: שאיפת שלטון"),
    ("פרידה מיונתן — אהבת נפש", "Wicked: \"For Good\" · Merrily We Roll Along: \"Old Friends\" · פנטום: \"All I Ask of You\""),
    ("מלחמות ועלייה", "המילטון: \"Yorktown\" · ספרטקוס · בן-הור · ממלכת שמים"),
    ("הריקוד ומיכל — אקסטזה מול בושה", "Wicked: \"Defying Gravity\" · מלך ואני: \"Shall We Dance\" · שנות ה-50"),
    ("לא אתה תבנה — קבלת מגבלה", "המילטון: \"The Room Where It Happens\" · ריצ'רד ב' (הסרת כתר) · Pippin: \"Corner of the Sky\""),
    ("בת-שבע ואוריה — תאווה ופשע", "מקבת: תכנון הרצח · דון ג'ובאני: פיתוי ומוות · אותלו: קנאה ומזימה · המלט: אופליה"),
    ("אמנון ותמר — פשע ודממה", "ריגולטו: אונס גילדה · טיטוס אנדרוניקוס · Parade (1998): אשמה ובדידות"),
    ("מרד אבשלום — קרע אב-בן", "ליר המלך: קורדליה/גונריל · ליאון מלך: שיבת סימבה · המלט: נקמת לארטס · המילטון: \"The Room...\""),
    ("מות אבשלום — האבל הגדול", "המילטון: \"It's Quiet Uptown\" · ליר+קורדליה · Sweeney Todd: אובדן · Rent: \"Without You\""),
    ("שיבה לירושלים — תיקון חלקי", "המילטון: \"The Election of 1800\" · לה מיזרבל: ז'אן ולז'אן חוזר · Pippin: home"),
    ("הללויה + שיר — כפרה דרך שיר", "המילטון: \"Who Lives, Who Dies\" · לה מיזרבל: \"Do You Hear the People Sing\" · Wicked: \"For Good\" · אביזר פינאלה"),
    ("הצוואה — ירושה מורכבת", "המילטון: \"Legacy\" · הנרי ד' חלק ב' · המלט: \"The rest is silence\" · ליאון מלך אפילוג"),
]

# --- Biblical coverage table ---
doc.add_page_break()
add_par("נספח — כיסוי פרקי התנ\"ך", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_par("הטבלה שלהלן מציגה לאיזה פרק בתנ\"ך מתייחסת כל סצנה במחזמר.", space_after=6)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = "פרק תנ\"ך"
hdr[1].text = "שם הסצנה"
hdr[2].text = "חלק"
for ref, name, part in COVERAGE:
    row = tbl.add_row().cells
    row[0].text = ref
    row[1].text = name
    row[2].text = part

# --- Modern references table ---
doc.add_page_break()
add_par("נספח — מקבילות דרמטורגיות", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_par("מקבילות בתיאטרון ובקולנוע לכל נושא דרמטי מרכזי במחזמר.", space_after=6)

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
hdr2[0].text = "נושא"
hdr2[1].text = "מקבילות בתיאטרון ובקולנוע"
for topic, refs in MODERN_REFS:
    row = tbl2.add_row().cells
    row[0].text = topic
    row[1].text = refs

out = ROOT / "מזמור-לדוד-ספר-לקריאה.docx"
doc.save(out)
print(f"נשמר: {out}")
print(f"מספר פסקאות: {len(doc.paragraphs)}")
