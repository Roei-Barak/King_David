#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
מחלץ מלל נקי מהמחזה — דיאלוג + שירים בלבד.
ללא הוראות במאי, ללא מקורות, ללא כותרות-ביניים ארגוניות.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/home/user/King_David")
SCENES_DIR = ROOT / "scenes"
FONT = "Arial"

SCENE_ORDER = [
    "scene-01-05-full.md",
    "scene-jonathan-farewell.md",
    "scene-gath-madness.md",
    "scene-nob-massacre.md",
    "scene-04-the-cave.md",
    "scene-04b-wilderness-additions.md",
    "scene-cave-hill.md",
    "scene-nabal-abigail.md",
    "scene-05b-witch-endor.md",
    "scene-ziklag.md",
    "scene-philistines-reject.md",
    "scene-gilboa.md",
    "scene-06-lament.md",
    "scene-07-brothers-war.md",
    "scene-ishbosheth.md",
    "scene-jerusalem.md",
    "scene-08-dance.md",
    "scene-09-not-you.md",
    "scene-chr25-singers.md",
    "scene-10-mephibosheth.md",
    "scene-10b-ammon-war.md",
    "scene-bathsheba-uriah.md",
    "scene-bathhouse.md",
    "scene-13b-amnon-tamar.md",
    "scene-13e-absalom-return.md",
    "scene-13c-revolt.md",
    "scene-13d-absalom.md",
    "scene-13f-david-return.md",
    "scene-13g-rizpah.md",
    "scene-water-libation.md",
    "scene-14-hallelujah.md",
    "scene-22-23-david-song.md",
    "scene-15-testament.md",
    "scene-tzruya-david.md",
]

# ---- ניקוי טקסט ----

def clean_text(text):
    """מסיר מטא-אלמנטים מתוך שורה"""
    # הוראות במאי inline *(...)*
    text = re.sub(r'\s*\*\([^)]*\)\*', '', text)
    # מקורות *(שמ"א ...)* / *(תה' ...)* וכד'
    text = re.sub(r'\s*\*\([^)]*\)\*', '', text)
    # סוגריים עם מקור מקראי: (שמ"א כ, ד) / (תהילים כב) וכד'
    text = re.sub(r'\s*\([^)]*(?:שמ|תה|מל|דב|בר|שו|עמ|יר|יש|יח|זכ|דה|נח|יו|רו|אי|מש|קה|שה|יב|בם|עז|נח)\s*[״"\']*[א-ת]*[^)]*\)', '', text)
    # [לאמת...] [מ-...]
    text = re.sub(r'\[[^\]]*\]', '', text)
    # bold / italic markers
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\*)([^*\n]+)(?<!\*)\*(?!\*)', r'\1', text)
    # backtick
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # רווחים כפולים
    text = re.sub(r'  +', ' ', text)
    return text.strip()

def is_scene_title(line, level):
    """כותרת # ראשית = שם סצנה (רמה 1 בלבד)"""
    return level == 1

def is_sub_heading(line, level):
    """כותרת ##+ — רק אם היא שם חלק עלילתי (לא ארגוני)"""
    if level < 2:
        return False
    s = re.sub(r'^#+\s*', '', line).strip()
    # כותרות ארגוניות/מטא — לדלג
    skip_words = [
        'חלק', 'פרק', 'שלב', 'סיכום', 'מקורות', 'מבנה', 'הוראות', 'שכבה',
        'תפקיד', 'טבלה', 'דרמטי', 'מוזיקה', 'תאורה', 'מפרט', 'רקע',
        'composer', 'brief', 'director', 'notes', 'פתיחה', 'פתיח',
        'גשר', 'כיסוי', 'source', 'מקור', 'שורש', 'אקספוז',
        'הערה', 'ז\'אנר', 'beat', 'transition', 'מעבר',
    ]
    sl = s.lower()
    for w in skip_words:
        if w in sl:
            return False
    # כותרת שיר / מונולוג — כן
    for w in ['שיר', 'מזמור', 'מונולוג', 'קינה', 'תפילה', '🎵', '🎭', 'דואט', 'פזמון']:
        if w in s:
            return True
    # כותרת קצרה שמזכירה שם דמות / מקום — כן
    return False

def extract_dialogue(line):
    """מחלץ (שם_דובר, תוכן) משורת דיאלוג **שם**: "..."  """
    m = re.match(r'^\*\*([^*]+)\*\*\s*(?:\*\([^)]*\)\*)?\s*[:：]\s*(.*)', line)
    if m:
        speaker = m.group(1).strip()
        # דלג על "דובר" שהוא הוראת במאי
        if speaker.startswith('*(') or speaker.lower() in ['הוראות', 'notes']:
            return None
        content = clean_text(m.group(2))
        return speaker, content
    return None

def is_song_line(line):
    s = line.strip()
    if not s.startswith('>'):
        return False
    inner = s[1:].strip()
    # הוראת במאי בתוך >
    if re.match(r'^\*\([^)]*\)\*\s*$', inner):
        return False
    if not inner:
        return False
    return True

def skip_line(line):
    s = line.strip()
    if not s:
        return False
    if s.startswith('|'):          return True   # טבלה
    if s.startswith('```'):        return True   # קוד
    if re.match(r'^>\s*$', s):    return True   # > ריק
    # הוראת במאי טהורה
    if re.match(r'^\*\([^)]*\)\*\.?\s*$', s):  return True
    if re.match(r'^\*[^*].*[^*]\*\s*$', s) and not re.match(r'^\*\*', s): return True
    # bullet list
    if re.match(r'^[-*+]\s+\S', s) and not re.match(r'^\*\*[^*]+\*\*', s): return True
    # שורות מקור בלבד
    if re.match(r'^[>*\s]*\*?\((?:שמ|תה|מל|דב|בר|שו|עמ|יר|יש|יח|זכ|דה|נח|יו|רו|אי|מש|קה|שה|יב|בם|עז)', s): return True
    # ### עם תוכן מטא
    return False

# ---- document ----

doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.2)

style = doc.styles["Normal"]
style.font.name = FONT; style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:cs"), FONT)
style.paragraph_format.space_after = Pt(0)

def rtl(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))

def par(text, size=11, bold=False, italic=False, color=None,
        center=False, sb=0, sa=2, ls=1.45):
    if not text or not text.strip(): return
    p = doc.add_paragraph(); rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(size * ls)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.name = FONT; r.font._element.rPr.rFonts.set(qn("w:cs"), FONT)
    if color: r.font.color.rgb = color
    return p

def dlg(speaker, content):
    if not content: return
    p = doc.add_paragraph(); rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(11 * 1.45)
    r1 = p.add_run(speaker + ":  ")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.name = FONT; r1.font._element.rPr.rFonts.set(qn("w:cs"), FONT)
    r1.font.color.rgb = RGBColor(0x22, 0x22, 0x66)
    r2 = p.add_run(content)
    r2.font.size = Pt(11)
    r2.font.name = FONT; r2.font._element.rPr.rFonts.set(qn("w:cs"), FONT)

def song_line(text):
    par(text, size=11, italic=True, center=True, sa=1, color=RGBColor(0x2a,0x2a,0x55))

# שער
for _ in range(4): doc.add_paragraph()
par("מִזְמוֹר לְדָוִד", size=34, bold=True, center=True, sa=10)
par("מחזמר — מלל", size=15, italic=True, center=True, sa=4)
doc.add_page_break()

# ---- עיבוד סצנה ----

def render(path):
    lines = path.read_text(encoding="utf-8").split("\n")
    in_code = False
    skip_block = False   # אחרי כותרת מטא
    song_buf = []
    last_was_dialogue = False

    def flush_song():
        nonlocal song_buf
        if song_buf:
            doc.add_paragraph()
            for sl in song_buf: song_line(sl)
            doc.add_paragraph()
        song_buf = []

    for raw in lines:
        line = raw.rstrip()

        if line.strip().startswith("```"):
            in_code = not in_code; continue
        if in_code: continue

        if not line.strip():
            flush_song()
            continue

        # כותרות
        hm = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm:
            flush_song()
            level = len(hm.group(1))
            title_raw = hm.group(2).strip()
            title = clean_text(title_raw)

            if is_scene_title(line, level):
                # שם הסצנה — כותרת עמוד
                doc.add_page_break()
                par(title, size=16, bold=True, center=True, sb=8, sa=8)
                skip_block = False
            elif is_sub_heading(line, level):
                par(title, size=12, bold=True, sb=6, sa=3,
                    color=RGBColor(0x33,0x33,0x66))
                skip_block = False
            else:
                # כותרת ארגונית — דלג עליה ועל תוכנה
                skip_block = False   # לא בולעים תוכן אחריה, רק מדלגים על הכותרת עצמה
            continue

        if skip_block: continue
        if skip_line(line): continue

        # שיר
        if is_song_line(line):
            inner = clean_text(line.strip()[1:])
            if inner: song_buf.append(inner)
            continue

        # דיאלוג
        d = extract_dialogue(line)
        if d:
            flush_song()
            speaker, content = d
            if content:
                dlg(speaker, content)
                last_was_dialogue = True
            continue

        # המשך ציטוט בשורה חדשה (מרכאות פותחות)
        s = line.strip()
        if (s.startswith('"') or s.startswith('״') or s.startswith('"')) and last_was_dialogue:
            content = clean_text(s)
            if content:
                p = doc.add_paragraph(); rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = Pt(11*1.45)
                p.paragraph_format.left_indent = Inches(0.25)
                r = p.add_run(content)
                r.font.size = Pt(11); r.font.name = FONT
                r.font._element.rPr.rFonts.set(qn("w:cs"), FONT)
            continue

        last_was_dialogue = False

    flush_song()

# ריצה
for fname in SCENE_ORDER:
    fpath = SCENES_DIR / fname
    if not fpath.exists():
        print(f"  חסר: {fname}"); continue
    print(f"  + {fname}")
    render(fpath)

out = ROOT / "מזמור-לדוד-מלל-נקי.docx"
doc.save(out)
print(f"\nנשמר: {out}  ({len(doc.paragraphs)} פסקאות)")
