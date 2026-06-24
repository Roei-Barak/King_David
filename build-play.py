#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build-play.py — בונה את "המחזה המושלם": מהדורת קריאה נקייה של המחזמר.
כולל: כותרות תמונות, הוראות במאי אמיתיות, דיאלוג, מונולוגים, שירים, ציטוטים מקראיים.
מנקה: הערות עבודה [לאמת]/[מ-x], אזכורי קבצים, "ממתינה לאישור", תזכירי הפקה,
       טבלאות כיסוי/מקבילות, סימוני מקור-מדרש פנימיים.
פלט: מיושר לימין (RTL), טיפוגרפיה נוחה לקריאה.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/home/user/King_David")
SCENES = ROOT / "scenes"
SONGS = ROOT / "songs"
FONT = "David"          # גופן עברי קלאסי; Word ימיר אם חסר
FONT_CS = "David"

# ---- סדר קנוני של המחזה ----
PLAY = [
    ("PART", "מַעֲרָכָה א׳", "מִן הָרוֹעֶה אֶל הַמֶּלֶךְ"),
    ("SCENE", "prologue.md", None),
    ("SCENE", "scene-01-05-full.md", None),
    ("SCENE", "scene-jonathan-farewell.md", None),
    ("SCENE", "scene-gath-madness.md", None),
    ("SCENE", "scene-nob-massacre.md", None),
    ("SCENE", "scene-04-the-cave.md", None),
    ("SCENE", "scene-04b-wilderness-additions.md", None),
    ("SCENE", "scene-nabal-abigail.md", None),
    ("SCENE", "scene-cave-hill.md", None),
    ("SCENE", "scene-05b-witch-endor.md", None),
    ("SCENE", "scene-ziklag.md", None),
    ("SCENE", "scene-philistines-reject.md", None),
    ("SCENE", "scene-gilboa.md", None),
    ("PART", "מַעֲרָכָה ב׳", "הַמֶּלֶךְ"),
    ("SCENE", "scene-06-lament.md", None),
    ("SCENE", "scene-07-brothers-war.md", None),
    ("SCENE", "scene-ishbosheth.md", None),
    ("SCENE", "scene-jerusalem.md", None),
    ("SCENE", "scene-08-dance.md", None),
    ("SCENE", "scene-09-not-you.md", None),
    ("SCENE", "scene-chr25-singers.md", None),
    ("SCENE", "scene-10-mephibosheth.md", None),
    ("SCENE", "scene-10b-ammon-war.md", None),
    ("SCENE", "scene-bathsheba-uriah.md", None),
    ("SCENE", "scene-bathhouse.md", None),
    ("SCENE", "scene-13b-amnon-tamar.md", None),
    ("SCENE", "scene-13e-absalom-return.md", None),
    ("SCENE", "scene-13c-revolt.md", None),
    ("SCENE", "scene-13d-absalom.md", None),
    ("SCENE", "scene-13f-david-return.md", None),
    ("SCENE", "scene-sheva-bichri.md", None),
    ("SCENE", "scene-13g-rizpah.md", None),
    ("SCENE", "scene-water-libation.md", None),
    ("SCENE", "scene-14-hallelujah.md", None),
    ("SCENE", "scene-22-23-david-song.md", None),
    ("SCENE", "scene-15-testament.md", None),
    ("PART", "נִסְפָּח", "תֵּת-עֲלִילָה — בְּנֵי צְרוּיָה"),
    ("SCENE", "scene-tzruya-david.md", None),
]

# ---- זיהוי מטא ----

# מילות-מפתח שמסמנות הערת-עבודה (אם מופיעות בתוך *(...)* או > הוא נמחק)
META_KW = [
    "מדרש [מ", "משובץ מדרש", "מקור:", "אומת", "אומתו", "לאמת", "Sefaria",
    "scene-", "songs/", "concept/", "research/", "sources/", ".md",
    "ראה `", "ממתינה לאישור", "Google Drive", "יובא מ", "[מ-",
    "כיוון ביצוע", "הערת בימוי", "שכבת רקע", "סאבטקסט", "להכריע",
    "song-", "תזכיר", "מפת ", "audit", "כיסוי",
    "הנמשל", "הערה לבמאי", "הערת טון", "⚠️", "style-guide", "style_guide",
]

# כותרות-מטא — מדלגים על הכותרת ועל גופהּ עד לכותרת/מפריד הבא
META_HEADINGS = [
    "סיכום", "מקורות", "מקור ", "כיסוי", "הערות בימוי", "טבלת", "טבלה",
    "מבנה דרמטי", "שכבה", "תפקיד דרמטי", "מקבילה מודרנית", "מקבילות",
    "מפת", "נספח דרמטורגי", "הערות", "מה עובד", "מה חסר", "לבמאי",
    "ז'אנר", "מוזיקה ולחן", "תאורה", "כוריאוגרפיה", "מפרט", "audit",
    "דרמטורגי", "checklist", "צ'ק", "ביבליוגרפ",
    # מקטעי-מאמר של הבמאי/מחבר (אינם טקסט מבוצע)
    "במאי", "הבמאי", "נמשל", "סאבטקסט", "אל הקורא", "דבר אל",
    "מבנה הפרולוג", "על הבמה", "מילה אחרונה", "הקבלות", "ההקבלות",
    "צומת", "ממצאי", "רמז פותח", "פעם ועתה", "פעם לעתה",
]

# תוויות-הערה: אם הן "שם הדובר" — זו הערה, לא דמות
NOTE_LABELS = ["הערה", "הערת", "לבמאי", "כיוון ביצוע", "סאבטקסט", "תזכיר", "שורש"]

# תוויות-מבנה של שיר — לשמר (בית/פזמון/שיא...) ולא למחוק כהערה
SONG_STRUCT = ["בית", "פזמון", "שיא", "פתיחה", "פתיח", "סיום", "גשר", "מעבר",
               "רפריז", "מבוא", "בריג'", "בריג", "טור", "מקהלה", "סולו", "דואט"]
def is_song_label(s):
    s = s.strip()
    return bool(s) and any(s.startswith(w) for w in SONG_STRUCT) and len(s) <= 22

# שורת מראה-מקום עצמאית (שמואל ב' ז, ב / תהילים כג, א)
REF_BOOKS = ("שמואל", "מלכים", "תהילים", "תהלים", "מזמור", "דברי הימים",
             "בראשית", "שמות", "ויקרא", "במדבר", "דברים", "ישעיה", "ישעיהו",
             "ירמיה", "ירמיהו", "יחזקאל", "הושע", "יואל", "עמוס", "מיכה",
             "משלי", "איוב", "רות", "איכה", "קהלת", "אסתר", "דניאל",
             "עזרא", "נחמיה", "זכריה", "מלאכי", "חבקוק", "צפניה")
def is_ref_line(s):
    s = s.strip().strip("()").strip()
    if len(s) > 34:
        return False
    if not any(s.startswith(b) for b in REF_BOOKS):
        return False
    return bool(re.search(r"[,׳\"']|[א-ת]['\"]?\s*[,]", s)) or bool(re.search(r"\d", s))

# קיצורי ספרי מקרא לזיהוי ציטוט בתוך *(...)*
BOOK_RE = (r"שמ[\"׳']?[אב]|תה(?:ילים|לים)?|מל[\"׳']?[אב]|דה[\"׳']?[אב]|"
           r"דב(?:רים)?|בר(?:אשית)?|שמות|וי(?:קרא)?|במ(?:דבר)?|"
           r"יש(?:עיה)?|יר(?:מיה)?|יח(?:זקאל)?|הו(?:שע)?|יואל|עמ(?:וס)?|"
           r"עו(?:בדיה)?|יו(?:נה)?|מי(?:כה)?|נח(?:ום)?|חב(?:קוק)?|"
           r"צפ(?:ניה)?|חג(?:י)?|זכ(?:ריה)?|מל(?:אכי)?|"
           r"מש(?:לי)?|אי(?:וב)?|שה[\"׳']?ש|רו(?:ת)?|איכה|קה(?:לת)?|"
           r"אס(?:תר)?|דנ(?:יאל)?|עז(?:רא)?|נחמ(?:יה)?|דה[\"׳']?י")

def is_citation_paren(inner):
    """האם תוכן הסוגריים הוא ציטוט מקור (שמ"א כד, ז) ולא הוראת במאי"""
    s = inner.strip().lstrip("*").strip()
    return bool(re.match(rf"^(?:{BOOK_RE})\b", s)) and len(s) < 40

def has_meta_kw(s):
    return any(kw in s for kw in META_KW)

def strip_brackets(s):
    """מסיר [..] עריכה; משאיר תוויות-מבנה של שיר ([בית א'] → בית א')"""
    def repl(m):
        inner = m.group(1).strip()
        if is_song_label(inner):
            return inner            # שמר תווית-שיר כטקסט
        return ""                   # הערת-עריכה — תרד
    return re.sub(r"\[([^\]]*)\]", repl, s)

def tidy(s):
    """ניקוי בסיס: סוגריים-מרובעים, backticks, הדגשות, מונחי-במה, שאריות *"""
    s = strip_brackets(s)
    s = re.sub(r"`[^`]*`", "", s)
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
    s = re.sub(r"\bBeat\b", "מַהֲלָךְ", s)
    s = re.sub(r"\bblackout\b", "חֲשֵׁכָה", s, flags=re.I)
    s = re.sub(r"\bdrone\b", "רֶחֶשׁ מִתְמַשֵּׁךְ", s, flags=re.I)
    s = s.replace("*", "")
    s = re.sub(r"  +", " ", s)
    return s.strip()

def clean_inline(s):
    """מנקה סימוני markdown + הערות-עבודה, שומר על ציטוטי-מקרא *(..)*"""
    s = strip_brackets(s)
    s = re.sub(r"`[^`]*`", "", s)
    # טפל בכל *(...)*: שמור ציטוט, מחק הוראה-בתוך-דיבור והערות-עבודה
    def paren_repl(m):
        inner = m.group(1)
        if is_citation_paren(inner):
            return f" ({inner.strip().lstrip('*').strip()})"
        return ""   # הוראת במאי בתוך שורת דיבור — תרד
    s = re.sub(r"\*\(([^)]*)\)\*", paren_repl, s)
    return tidy(s)

# ---- מסמך ----
doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.15); sec.right_margin = Inches(1.15)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:cs"), FONT_CS)
style.paragraph_format.space_after = Pt(0)

# RTL ברמת המסמך (ברירת מחדל לכל פסקה חדשה)
def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi"); bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

def P(text, size=12, bold=False, italic=False, color=None, align="right",
      sb=0, sa=4, ls=1.4, indent=0.0, first=0.0):
    if text is None or text == "":
        return None
    p = doc.add_paragraph(); set_rtl(p)
    p.alignment = {"right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    pf = p.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing = Pt(size * ls)
    if indent: pf.right_indent = Inches(indent)
    if first:  pf.first_line_indent = Inches(first)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.name = FONT
    r.font._element.rPr.rFonts.set(qn("w:cs"), FONT_CS)
    if color: r.font.color.rgb = color
    return p

def speaker_par(name, action, content):
    """שורת דיאלוג: שם (בולט) + [הוראה קצרה איטלית] + דיבור"""
    p = doc.add_paragraph(); set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(3); pf.space_after = Pt(3)
    pf.line_spacing = Pt(12 * 1.4)
    rn = p.add_run(name + ": ")
    rn.bold = True; rn.font.size = Pt(12)
    rn.font.name = FONT; rn.font._element.rPr.rFonts.set(qn("w:cs"), FONT_CS)
    rn.font.color.rgb = RGBColor(0x1a, 0x2a, 0x6b)
    if action:
        ra = p.add_run(f"({action}) ")
        ra.italic = True; ra.font.size = Pt(10.5)
        ra.font.name = FONT; ra.font._element.rPr.rFonts.set(qn("w:cs"), FONT_CS)
        ra.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    rc = p.add_run(content)
    rc.font.size = Pt(12)
    rc.font.name = FONT; rc.font._element.rPr.rFonts.set(qn("w:cs"), FONT_CS)
    return p

GRAY = RGBColor(0x70, 0x70, 0x70)
SONGCLR = RGBColor(0x23, 0x23, 0x5c)
TITLECLR = RGBColor(0x10, 0x10, 0x10)

# ---- עיבוד סצנה ----
def render_scene(text):
    lines = text.split("\n")
    in_code = False
    skip_meta = False         # בתוך מקטע-מטא (כותרת + גוף)
    skip_meta_level = 0       # הרמה שבה התחיל מקטע-המטא
    song_block = False        # רצף שורות-שיר ב->
    scene_title_done = False
    i = 0
    while i < len(lines):
        raw = lines[i]; i += 1
        line = raw.rstrip()

        if line.strip().startswith("```"):
            in_code = not in_code; continue
        if in_code:
            continue

        # שורה ריקה
        if not line.strip():
            song_block = False
            continue

        # טבלאות — תמיד מטא
        if line.strip().startswith("|"):
            continue

        # מפריד --- => מסיים מקטע-מטא ומצייר קו עדין
        if re.match(r"^\s*(---+|⸻|\*\s*\*\s*\*)\s*$", line):
            if skip_meta:
                skip_meta = False
                continue              # אל תצייר קו בתוך/בסיום מקטע-מטא
            P("◆", align="center", size=10, color=RGBColor(0xb0,0xb0,0xb0), sb=6, sa=6)
            continue

        # כותרות
        hm = re.match(r"^(#{1,6})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1))
            raw_title = hm.group(2).strip()
            is_meta = any(w in raw_title for w in META_HEADINGS) or has_meta_kw(raw_title)
            if is_meta:
                skip_meta = True
                skip_meta_level = level
                continue
            # כותרת רגילה — אך אם אנו בתוך מקטע-מטא וזו תת-כותרת, המשך לדלג
            if skip_meta and level > skip_meta_level:
                continue
            skip_meta = False
            title = clean_inline(raw_title)
            if not title:
                continue
            if level == 1 and not scene_title_done:
                P(title, size=19, bold=True, align="center", sb=4, sa=10,
                  color=TITLECLR)
                scene_title_done = True
            elif level == 1:
                P(title, size=16, bold=True, align="center", sb=14, sa=8, color=TITLECLR)
            elif level == 2:
                P(title, size=13.5, bold=True, sb=12, sa=5,
                  color=RGBColor(0x2a,0x2a,0x55))
            else:
                P(title, size=12, bold=True, sb=8, sa=3,
                  color=RGBColor(0x44,0x44,0x66))
            continue

        if skip_meta:
            continue

        # בלוק ציטוט > — שיר/פיוט או הערה
        if line.strip().startswith(">"):
            inner = line.strip()[1:].strip()
            if not inner:
                continue
            # הערת-עבודה/הפקה => דלג
            if inner.startswith("**") or has_meta_kw(inner) or inner.startswith("הערה"):
                continue
            txt = clean_inline(inner)
            if txt:
                P(txt, size=11.5, italic=True, align="center", color=SONGCLR,
                  sb=1, sa=1, ls=1.5)
            continue

        # רשימות תבליט — כמעט תמיד הערות; דלג
        if re.match(r"^\s*[-*+]\s+", line) and not re.match(r"^\s*\*\*", line):
            continue

        # הוראת במאי שורה-שלמה *(...)*
        sd = re.match(r"^\s*\*\((.+)\)\*\s*$", line)
        if sd:
            inner = sd.group(1)
            if has_meta_kw(inner):       # הוראה שהיא בעצם הערת-מקור => דלג
                continue
            body = tidy(inner)
            if body:
                P(body, size=10.5, italic=True, color=GRAY, align="right",
                  sb=2, sa=2, indent=0.15)
            continue

        # שורת דיאלוג: **שם** *(הוראה)*: דיבור
        dm = re.match(r"^\*\*([^*]+)\*\*\s*(?:\*\(([^)]*)\)\*)?\s*[:：]\s*(.*)$", line)
        if dm:
            name = dm.group(1).strip()
            # "שם דובר" שהוא בעצם תווית-הערה => דלג
            if any(nl in name for nl in NOTE_LABELS):
                continue
            action = (dm.group(2) or "").strip()
            content = clean_inline(dm.group(3))
            if action and has_meta_kw(action):
                action = ""
            # ננקה הוראה ארוכה מדי לשורית
            if action:
                action = clean_inline("*(" + action + ")*").strip("() ")
            if content:
                speaker_par(name, action, content)
            else:
                # שם בלבד (דיבור בשורות הבאות)
                speaker_par(name, action, "")
            continue

        # שורת-טקסט פשוטה שהיא בעצם הערת-עבודה => דלג
        if has_meta_kw(line) or re.match(r"^\s*גִּ?רְ?סָ?ה?\s*רִ?אשׁ?וֹ?נָ?ה?", line) \
           or line.strip().startswith("גרסה ראשונה"):
            continue

        # שורת המשך/מונולוג (פסוק חופשי או המשך דיבור)
        txt = clean_inline(line)
        if txt:
            # תווית-מבנה של שיר (בית/פזמון/שיא) — תצוגה מובחנת
            if is_song_label(txt):
                P(txt, size=10.5, bold=True, align="center", color=GRAY,
                  sb=5, sa=1)
            elif is_ref_line(txt):
                # מראה-מקום עצמאי — קטן ואפור, נסוג מן הטקסט
                P("(" + txt.strip("()").strip() + ")", size=9.5, color=GRAY,
                  align="right", sb=0, sa=2, indent=0.25)
            else:
                # שורות מונולוג/פסוק — הזחה קלה, שמירת שבירת השורה
                P(txt, size=12, align="right", sb=0, sa=2, ls=1.4, indent=0.25)

    return

# ---- שער ----
for _ in range(5): doc.add_paragraph()
P("מִזְמוֹר לְדָוִד", size=40, bold=True, align="center", sa=10, color=TITLECLR)
P("הַצַּדִּיק שֶׁכּוֹתֵב שִׁירִים", size=20, italic=True, align="center", sa=22,
  color=RGBColor(0x33,0x33,0x55))
P("מַחֲזֶמֶר בָּמָה עַל חַיֵּי דָּוִד הַמֶּלֶךְ", size=14, align="center", sa=4,
  color=RGBColor(0x55,0x55,0x55))
P("אֱמוּנָה דֶּרֶךְ יִסּוּרִים — כָּל עֲצִירַת יָד מוֹלֶדֶת כִּנּוֹר", size=12.5,
  italic=True, align="center", color=RGBColor(0x77,0x77,0x77))
doc.add_page_break()

# ---- בנייה ----
count_scenes = 0
for kind, a, b in PLAY:
    if kind == "PART":
        doc.add_page_break()
        for _ in range(6): doc.add_paragraph()
        P(a, size=30, bold=True, align="center", sa=8, color=TITLECLR)
        if b:
            P(b, size=17, italic=True, align="center", color=RGBColor(0x55,0x55,0x55))
        continue
    fpath = SCENES / a
    if not fpath.exists():
        print(f"  חסר: {a}")
        continue
    doc.add_page_break()
    print(f"  + {a}")
    render_scene(fpath.read_text(encoding="utf-8"))
    count_scenes += 1

out = ROOT / "מזמור-לדוד-המחזה.docx"
doc.save(out)
print(f"\nנשמר: {out}")
print(f"תמונות: {count_scenes} · פסקאות: {len(doc.paragraphs)}")
