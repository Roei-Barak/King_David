#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ממיר את קבצי הסצנות של "מזמור לדוד" לסיפור פרוזה ספרותי.

כללי ההמרה:
- דיאלוג  **שם**: "..." → "..." אמר/ענה/פנה שם
- הוראות במאי *(...)* → פסקת תיאור
- מונולוג שייקספירי → קטע פנימי "בלבו חשב:"
- מונולוג תהילים (דוד) → שיר מצוטט מוקף גרש
- שיר → שיר מצוטט עם יסורת
- כותרות ## → כותרת פרק
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/home/user/King_David")
SCENES_DIR = ROOT / "scenes"
SONGS_DIR = ROOT / "songs"

# סדר הסצנות — זהה ל-build-book.py
BOOK = [
    ("part", "מַעֲרָכָה א — מִן הָרוֹעֶה לַמֶּלֶךְ"),
    ("scene", "prologue.md"),
    ("scene", "scene-01-05-full.md"),
    ("part", "הרחבות המדבר"),
    ("scene", "scene-jonathan-farewell.md"),
    ("scene", "scene-gath-madness.md"),
    ("scene", "scene-nob-massacre.md"),
    ("scene", "scene-04-the-cave.md"),
    ("scene", "scene-04b-wilderness-additions.md"),
    ("scene", "scene-nabal-abigail.md"),
    ("scene", "scene-cave-hill.md"),
    ("scene", "scene-05b-witch-endor.md"),
    ("part", "מַעֲרָכָה ב — הַמֶּלֶךְ"),
    ("scene", "scene-06-lament.md"),
    ("scene", "scene-07-brothers-war.md"),
    ("scene", "scene-ishbosheth.md"),
    ("scene", "scene-jerusalem.md"),
    ("scene", "scene-08-dance.md"),
    ("scene", "scene-09-not-you.md"),
    ("scene", "scene-10-mephibosheth.md"),
    ("scene", "scene-10b-ammon-war.md"),
    ("scene", "scene-bathsheba-uriah.md"),
    ("scene", "scene-13b-amnon-tamar.md"),
    ("scene", "scene-13e-absalom-return.md"),
    ("scene", "scene-13c-revolt.md"),
    ("scene", "scene-13d-absalom.md"),
    ("scene", "scene-13f-david-return.md"),
    ("scene", "scene-13g-rizpah.md"),
    ("scene", "scene-14-hallelujah.md"),
    ("scene", "scene-15-testament.md"),
    ("scene", "scene-tzruya-david.md"),
]

FONT = "Arial"

# פועלי אמירה לסירוגין (כדי לא לחזור על "אמר")
SPEECH_VERBS = {
    "default": ["אמר", "ענה", "פנה", "קרא", "אמר"],
    "quiet":   ["לחש", "אמר בשקט", "מלמל"],
    "shout":   ["קרא", "צעק", "זעק"],
    "narrator":["מספר המחבר"],
}

# שמות לעברית — ממפה כתיב ניקוד/בלי ניקוד
NAME_MAP = {
    "דָּוִד": "דוד", "דוד": "דוד",
    "שָׁאוּל": "שאול", "שאול": "שאול",
    "יְהוֹנָתָן": "יונתן", "יונתן": "יונתן",
    "מִיכַל": "מיכל", "מיכל": "מיכל",
    "אֲבִיגַיִל": "אביגיל", "אביגיל": "אביגיל",
    "יוֹאָב": "יואב", "יואב": "יואב",
    "נָתָן": "נתן", "נתן": "נתן",
    "אַבְשָׁלוֹם": "אבשלום", "אבשלום": "אבשלום",
    "בַּת-שֶׁבַע": "בת-שבע", "בת-שבע": "בת-שבע",
    "אֲחִיתֹפֶל": "אחיתופל", "אחיתופל": "אחיתופל",
    "אֶלְיָאב": "אליאב", "אליאב": "אליאב",
    "אֲבִישַׁי": "אבישי", "אבישי": "אבישי",
    "אַבְנֵר": "אבנר", "אבנר": "אבנר",
    "שְׁמוּאֵל": "שמואל", "שמואל": "שמואל",
    "יִשַׁי": "ישי", "ישי": "ישי",
    "אוּרִיָּה": "אוריה", "אוריה": "אוריה",
    "שְׁלֹמֹה": "שלמה", "שלמה": "שלמה",
    "מְסַפֵּר": "המספר", "קוֹל": "קול",
    "אַרְוְנָה": "ארוונה",
    "נָבָל": "נבל",
    "רִצְפָּה": "רצפה",
    "אֲחִימֶלֶךְ": "אחימלך",
    "דּוֹאֵג": "דואג",
    "אֵלִיָּהוּ": "אליהו",
    "צְרוּיָה": "צרויה",
    "שַׁלִּיחַ": "שליח",
    "חוּשַׁי": "חושי",
    "שִׁמְעִי": "שמעי",
    "מְפִיבֹשֶׁת": "מפיבושת",
    "גָּד": "גד",
    "חִירָם": "חירם",
}

def strip_nikud(s):
    """הסרת ניקוד ומפות לשם פשוט."""
    return NAME_MAP.get(s.strip(), s.strip())

def clean_md_inline(s):
    """הסרת סימוני Markdown inline."""
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = re.sub(r"\[לאמת[^\]]*\]", "", s)
    s = re.sub(r"\*\([^)]+\)\*", "", s)   # *(source)* annotations
    s = s.strip()
    return s

def extract_stage(s):
    """מחלץ הוראות במאי מתוך *(...)* — מחזיר טקסט ללא הסוגריים."""
    matches = re.findall(r"\*\(([^)]+)\)\*", s)
    return " ".join(matches)

def remove_stage(s):
    """מוחק *(...)* מטקסט."""
    return re.sub(r"\*\([^)]*\)\*", "", s).strip()

# ===== Word document helpers =====

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:cs"), FONT)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)

def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_par(text, size=12, bold=False, italic=False, color=None,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=6,
            line_spacing=1.4, center=False):
    if not text.strip():
        return
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * line_spacing)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run.font._element.rPr.rFonts.set(qn("w:cs"), FONT)
    if color:
        run.font.color.rgb = color
    return p

def add_separator():
    add_par("* * *", center=True, size=11, color=RGBColor(0x99,0x99,0x99), space_before=6, space_after=6)

def add_poem_line(text, size=11):
    """שורת שיר — מרכז, אותיות מוטות."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(size * 1.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = True
    run.font.name = FONT
    run.font._element.rPr.rFonts.set(qn("w:cs"), FONT)
    run.font.color.rgb = RGBColor(0x22,0x22,0x55)

# ===== שער =====
for _ in range(4): doc.add_paragraph()
add_par("מִזְמוֹר לְדָוִד", size=36, bold=True, center=True, space_after=10)
add_par("סִפּוּר חַיָּיו שֶׁל הַמֶּלֶךְ", size=18, italic=True, center=True, space_after=8)
add_par("מֵהַשָּׂדֶה לַבֵּית-הַמִּקְדָּשׁ — אֱמוּנָה דֶּרֶךְ יְסוּרִים", size=13, center=True, space_after=4)
doc.add_page_break()

add_par("לַקּוֹרֵא", size=16, bold=True, space_after=6)
add_par(
    "מה שלפניך הוא הסיפור כפי שנוצר: דרמה מקראית ועברית על חייו של דוד המלך — "
    "מן הנער הרועה צאן ועד המלך הזקן המוסר את ממלכתו לבנו שלמה. "
    "הסיפור מבוסס על ספרי שמואל, מלכים ודברי הימים, ומשלב מדרשים ופרשנויות חז\"ל. "
    "שירי תהילים — קולו הפנימי של דוד — מצוטטים בתוך הסיפור כפי שנולדו: מתוך הרגע.", space_after=12)
doc.add_page_break()

# ===== מנגנון המרת סצנה לפרוזה =====

def classify_action(stage_text):
    """מחליט מה לעשות עם הוראת במאי."""
    if not stage_text:
        return None
    low = stage_text.lower()
    skip_patterns = [
        "הוראות במאי", "הערת", "אפשרות", "אופציה", "להכריע",
        "סאבטקסט", "שכבת", "ניגוד מכוון", "מדרש", "[מ-", "ראה `",
        "note:", "note —", "מדרש [",
    ]
    for pat in skip_patterns:
        if pat in stage_text:
            return None
    return stage_text

# סעיפים שכולם הוראות במאי / מטא — לדלג עליהם לחלוטין
META_SECTION_PATTERNS = [
    "סאבטקסט", "הוראות במאי", "הערות בימוי", "הערת בימוי",
    "דמויות בסצנה", "דמויות:", "דמויות —",
    "אימות מקורות", "מראי-מקום", "מראה מקום",
    "שכבות הסצנה", "שכבת המלך", "שכבת",
    "חוט הזהב", "ממצאי", "מבנה הפרולוג",
    "ההקבלות", "הצומת המותזת",
    "מילה אחרונה לבמאי", "דבר אל הקורא ואל הבמאי",
    "הנמשל של הבמאי", "הוראה למלחין",
    "רמז פותח", "הרחבות:", "אופציית במאי",
    "נוכחות שקטה", "הערת מסגרת",
]

def is_meta_heading(txt):
    for pat in META_SECTION_PATTERNS:
        if pat in txt:
            return True
    return False

def prose_from_scene(md_text):
    """ממיר טקסט Markdown של סצנה לרשימת אובייקטי פרוזה."""
    lines = md_text.split("\n")
    result = []   # list of (type, text) — type: heading/prose/dialogue/poem/stage

    in_monologue = False
    in_song = False
    skip_meta_section = False   # מדלגים על קטע מטא
    monologue_buffer = []
    song_buffer = []
    current_speaker = None
    prev_verb_idx = 0
    verbs = ["אמר", "ענה", "פנה", "אמר", "השיב", "קרא", "דיבר"]

    def flush_monologue():
        nonlocal monologue_buffer, in_monologue
        if monologue_buffer:
            result.append(("monologue", "\n".join(monologue_buffer), current_speaker))
        monologue_buffer = []
        in_monologue = False

    def flush_song():
        nonlocal song_buffer, in_song
        if song_buffer:
            result.append(("song", "\n".join(song_buffer)))
        song_buffer = []
        in_song = False

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # --- מפריד ---
        if re.match(r"^\s*---+\s*$", line):
            flush_monologue(); flush_song()
            skip_meta_section = False
            result.append(("separator", ""))
            i += 1
            continue

        # --- כותרות ---
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            flush_monologue(); flush_song()
            level = len(m.group(1))
            txt = clean_md_inline(m.group(2))
            txt = re.sub(r"[🎭🎼🎵🎶⚔️🙏✡️]+", "", txt).strip()
            txt = re.sub(r"\*\([^)]*\)\*", "", txt).strip()
            # בדוק אם זה כותרת של קטע מטא
            if is_meta_heading(txt):
                skip_meta_section = True
                i += 1
                continue
            # כותרת רגילה — מאפסת skip
            skip_meta_section = False
            result.append(("heading", txt, level))
            i += 1
            continue

        # --- אם אנחנו בקטע מטא — דלג ---
        if skip_meta_section:
            i += 1
            continue

        # --- ציטוט/בלוק > ---
        if line.strip().startswith(">"):
            flush_monologue(); flush_song()
            txt = clean_md_inline(line.strip()[1:].strip())
            skip_content = ["הוראות", "הערת", "מדרש", "ראה `", "שכבת", "[מ-",
                            "מקור:", "זמן:", "שיר:", "סטטוס:", "מערכה:", "דמויות:"]
            if any(p in txt for p in skip_content):
                i += 1
                continue
            result.append(("quote", txt))
            i += 1
            continue

        # --- שורת טבלה ---
        if line.strip().startswith("|"):
            i += 1
            continue

        # --- שורת רשימה עם מטא-תוכן ---
        if re.match(r"^\s*[-*]\s+", line):
            txt = clean_md_inline(re.sub(r"^\s*[-*]\s+", "", line))
            meta_list = ["הוראות", "הערת", "מדרש", "במאי", "תאורה", "מוזיקה",
                         "שכבת", "ניגוד", "אופציה", "חוט", "ראה `"]
            if any(p in txt for p in meta_list):
                i += 1
                continue
            result.append(("prose", txt))
            i += 1
            continue

        # --- שורה ריקה ---
        if not line.strip():
            flush_monologue()
            i += 1
            continue

        # --- סימון מונולוג/שיר בכותרת משנה ---
        if re.search(r"מונולוג|שיר:|שיר #|🎭|🎼", line):
            flush_monologue(); flush_song()
            if "שיר" in line and "מונולוג" not in line:
                in_song = True
            else:
                in_monologue = True
            i += 1
            continue

        # --- דיאלוג: **שם** *(action)*: "..." ---
        # פורמט: **NAME** *(opt-action)*: "text"
        dialog_match = re.match(
            r'^\*\*([^*]+)\*\*\s*(?:\*\(([^)]*)\)\*)?\s*[:]\s*["״“”](.+)', line
        )
        if not dialog_match:
            # ניסיון נוסף — שם בלי גרשיים ישירים בשורה
            dialog_match = re.match(r"^\*\*([^*]+)\*\*\s*(?:\*\(([^)]*)\)\*)?\s*[:：]\s*(.*)", line)

        if dialog_match:
            flush_song()
            if in_monologue:
                # המשך מונולוג — אוסף
                txt = clean_md_inline(dialog_match.group(3))
                if txt:
                    monologue_buffer.append(txt)
                i += 1
                continue

            flush_monologue()
            raw_name = dialog_match.group(1)
            raw_action = dialog_match.group(2) or ""
            speech_raw = dialog_match.group(3)

            speaker = strip_nikud(raw_name)
            action = classify_action(clean_md_inline(raw_action))
            speech = clean_md_inline(speech_raw).strip('"\'״"').strip()

            # בדוק אם השורה הבאה המשך דיאלוג (מרכאות לא נסגרו)
            j = i + 1
            while j < len(lines):
                nxt = lines[j].rstrip()
                if not nxt.strip():
                    break
                if re.match(r"^\*\*", nxt) or re.match(r"^\s*---", nxt) or re.match(r"^#", nxt):
                    break
                if nxt.strip().startswith("*(") or nxt.strip().startswith("**("):
                    extra_action = extract_stage(nxt)
                    if extra_action:
                        action = (action + " " + extra_action).strip() if action else extra_action
                    j += 1
                    continue
                extra = clean_md_inline(nxt).strip('"\'״"').strip()
                if extra:
                    speech += " " + extra
                j += 1

            i = j

            # בחר פועל אמירה
            nonlocal_verb = verbs[prev_verb_idx % len(verbs)]
            prev_verb_idx += 1
            current_speaker = speaker

            result.append(("dialogue", speaker, speech, action, nonlocal_verb))
            continue

        # --- פסקת תיאור רגיל / הוראת במאי בשורה רגילה ---
        # שורה שמתחילה ב-*(  — הוראת במאי
        stage_match = re.match(r"^\*\((.+)\)\*\s*$", line)
        if stage_match:
            flush_monologue(); flush_song()
            txt = classify_action(clean_md_inline(stage_match.group(1)))
            if txt:
                result.append(("stage", txt))
            i += 1
            continue

        # --- שורה במצב מונולוג ---
        if in_monologue:
            txt = clean_md_inline(line)
            # מדלגים על שורות שהן source citations
            if re.match(r"^\*\(שמ|^\*\(תה|^\*\(מל|^\*\(בר|^\*\(דב", line):
                i += 1
                continue
            if txt:
                monologue_buffer.append(txt)
            i += 1
            continue

        # --- שורה במצב שיר ---
        if in_song:
            txt = clean_md_inline(line)
            if txt and not txt.startswith("|"):
                song_buffer.append(txt)
            i += 1
            continue

        # --- שורה רגילה — מתארת / נרטיב ---
        txt = clean_md_inline(line)
        if txt and not re.match(r"^\s*\|", line):
            # בדוק אם זה stage direction שלא תפסנו
            if line.strip().startswith("*(") and line.strip().endswith(")*"):
                inner = line.strip()[2:-2]
                classified = classify_action(clean_md_inline(inner))
                if classified:
                    flush_monologue()
                    result.append(("stage", classified))
            else:
                result.append(("prose", txt))
        i += 1

    flush_monologue()
    flush_song()
    return result


def render_to_doc(items):
    """כותב רשימת אובייקטים ל-Word document."""
    for item in items:
        kind = item[0]

        if kind == "separator":
            add_separator()

        elif kind == "heading":
            txt, level = item[1], item[2]
            if not txt:
                continue
            if level == 1:
                doc.add_page_break()
                add_par(txt, size=20, bold=True, center=True, space_before=8, space_after=8)
            elif level == 2:
                add_par(txt, size=15, bold=True, space_before=10, space_after=5)
            elif level == 3:
                add_par(txt, size=13, bold=True, italic=True,
                        color=RGBColor(0x33,0x33,0x66), space_before=6, space_after=4)
            else:
                add_par(txt, size=12, bold=True, space_before=4, space_after=3)

        elif kind == "prose":
            txt = item[1]
            if not txt:
                continue
            add_par(txt, size=12, line_spacing=1.5, space_after=6)

        elif kind == "stage":
            txt = item[1]
            if not txt:
                continue
            # הוראות במאי → פסקת תיאור נרטיבי
            add_par(txt, size=11.5, italic=True,
                    color=RGBColor(0x33,0x33,0x44), space_before=2, space_after=4)

        elif kind == "dialogue":
            _, speaker, speech, action, verb = item
            if not speech:
                continue
            # בנה משפט: "הדיבור," אמר הדמות, בפועל/ביתה.
            if action:
                action_clean = action.rstrip(".")
                prose = f"«{speech}» — {verb} {speaker}, {action_clean}."
            else:
                prose = f"«{speech}» — {verb} {speaker}."
            add_par(prose, size=12, line_spacing=1.5, space_after=6)

        elif kind == "monologue":
            _, text, speaker = item
            lines_m = [l for l in text.split("\n") if l.strip()]
            if not lines_m:
                continue
            if speaker:
                add_par(f"בלבו הרהר {speaker}:", size=11.5, italic=True,
                        color=RGBColor(0x22,0x22,0x55), space_before=4, space_after=2)
            for l in lines_m:
                add_poem_line(l.strip())
            doc.add_paragraph()  # רווח אחרי

        elif kind == "song":
            lines_s = [l for l in item[1].split("\n") if l.strip()]
            if not lines_s:
                continue
            add_par("—", center=True, size=10, color=RGBColor(0x99,0x99,0x99))
            for l in lines_s:
                add_poem_line(l.strip())
            add_par("—", center=True, size=10, color=RGBColor(0x99,0x99,0x99))

        elif kind == "quote":
            txt = item[1]
            if txt:
                add_par(f"״{txt}״", size=11, italic=True,
                        color=RGBColor(0x44,0x44,0x44), space_before=2, space_after=2)


# ===== בנייה =====
print("מתחיל בניית ספר הפרוזה...")

scene_count = 0
for entry in BOOK:
    kind = entry[0]
    if kind == "part":
        doc.add_page_break()
        add_par(entry[1], size=22, bold=True, center=True, space_before=20, space_after=10)
        add_par("⸻", center=True, color=RGBColor(0x88,0x88,0x88), space_after=16)
        continue

    fname = entry[1]
    # בדוק אם בסצנות או בשירים
    fpath = SCENES_DIR / fname
    if not fpath.exists():
        fpath = SONGS_DIR / fname
    if not fpath.exists():
        print(f"  דולג (לא קיים): {fname}")
        continue

    print(f"  מעבד: {fname}")
    md = fpath.read_text(encoding="utf-8")
    items = prose_from_scene(md)
    render_to_doc(items)
    scene_count += 1

out = ROOT / "מזמור-לדוד-סיפור.docx"
doc.save(out)
print(f"\nנשמר: {out}")
print(f"סצנות שעובדו: {scene_count}")
print(f"מספר פסקאות: {len(doc.paragraphs)}")
